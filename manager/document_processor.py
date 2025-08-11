# Copyright (c) 2025 Tsutomu FUNADA
# This software is licensed under the MIT License for non-commercial use.

import tempfile
from io import BytesIO
from typing import Optional

import fitz  # PDF (fitz == PyMuPDF)
from docx import Document  # Microsoft Word
from ebooklib import epub  # EPUB

from config.types import DOCUMENT_CONVERTIBLE_FORMATS, DOCUMENT_FILETYPE_MAP


class DocumentProcessor:
    """Handles document format conversions (PDF, EPUB, TXT, DOCX)."""

    def convert_document(
        self, format: str, content: bytes, mimetype: str
    ) -> Optional[bytes]:
        """
        Converts a document to the specified format.

        Args:
            format (str): Target format ('txt', 'pdf', 'docx', 'epub').
            content (bytes): Raw document content.
            mimetype (str): MIME type of the document.

        Returns:
            Optional[bytes]: Converted document content or None if unsupported format.
        """
        if format not in DOCUMENT_CONVERTIBLE_FORMATS or not mimetype:
            raise ValueError(f"Unsupported conversion: '{mimetype}' to '{format}'.")

        converters = {
            "txt": self.convert_to_text,
            "pdf": self.convert_to_pdf,
            "docx": self.convert_to_docx,
            "epub": self.convert_to_epub,
        }

        convert_func = converters.get(format)
        return convert_func(content, mimetype) if convert_func else None

    def convert_to_text(self, content: bytes, mimetype: str) -> bytes:
        """
        Converts supported document formats into plain text using fitz (PyMuPDF).

        Args:
            content (bytes): Raw binary content of the document.
            mimetype (str): MIME type of the document.

        Returns:
            bytes: Converted text content in UTF-8 encoding.
        """
        if mimetype == "text/plain":
            return content  # No conversion needed

        elif mimetype in DOCUMENT_FILETYPE_MAP:
            doc = fitz.open(
                stream=BytesIO(content), filetype=DOCUMENT_FILETYPE_MAP[mimetype]
            )
            text_content = "\n".join(page.get_text("text") for page in doc.pages())
            return text_content.encode("utf-8")

        else:
            raise ValueError(f"Unsupported format: {mimetype}")

    def convert_to_pdf(self, content: bytes, mimetype: str) -> bytes:
        """
        Converts supported document formats into PDF using PyMuPDF.

        Args:
            content (bytes): Raw binary content of the document.
            mimetype (str): MIME type of the document.

        Returns:
            bytes: Converted PDF content.
        """
        if mimetype == "application/pdf":
            return content  # No conversion needed

        elif mimetype in DOCUMENT_FILETYPE_MAP:
            doc = fitz.open(
                stream=BytesIO(content), filetype=DOCUMENT_FILETYPE_MAP[mimetype]
            )
            pdf_bytes = doc.convert_to_pdf()
            return pdf_bytes

        else:
            raise ValueError(f"Unsupported format: {mimetype}")

    def convert_to_epub(self, content: bytes, mimetype: str) -> bytes:
        """
        Converts the given document content into EPUB format.

        If the document is already EPUB, returns it unchanged.
        For other supported types, first converts to plain text and then
        creates an EPUB with a single chapter containing the text.

        Args:
            content (bytes): Raw binary content of the document.
            mimetype (str): MIME type of the document.

        Returns:
            bytes: Converted EPUB content.
        """
        if mimetype == "application/epub+zip":
            return content  # 既に EPUB の場合はそのまま返す

        # 他の形式の場合は、まずテキストに変換する
        plain_text = self.convert_to_text(content, mimetype).decode("utf-8")

        # EPUB のビルド
        book = epub.EpubBook()
        book.set_identifier("id123456")
        book.set_title("Converted Book")
        book.set_language("en")
        book.add_author("Unknown")

        # シンプルな1章に全テキストを収める
        chapter = epub.EpubHtml(title="Chapter 1", file_name="chap_01.xhtml", lang="en")
        # 改行ごとに<p>タグで囲む（必要に応じてさらに細かい解析が可能）
        chapter.content = (
            "<html><body><p>"
            + plain_text.replace("\n", "</p><p>")
            + "</p></body></html>"
        )
        book.add_item(chapter)

        # 目次 (TOC) やナビゲーション追加
        book.toc = [epub.Link("chap_01.xhtml", "Chapter 1", "chap1")]
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        # CSS を追加（必要に応じて調整可能）
        style = "BODY { font-family: Times, serif; }".encode("utf-8")
        nav_css = epub.EpubItem(
            uid="style_nav",
            file_name="style/nav.css",
            media_type="text/css",
            content=style,
        )
        book.add_item(nav_css)

        book.spine = ["nav", chapter]

        # epub.write_epub はファイル名を受け取るため、一時ファイルに書き出してバイト列を返す
        with tempfile.NamedTemporaryFile(suffix=".epub", delete=False) as tmp:
            epub.write_epub(tmp.name, book)
            tmp.seek(0)
            epub_data = tmp.read()

        return epub_data

    def convert_to_docx(self, content: bytes, mimetype: str) -> bytes:
        """
        Converts supported document formats into DOCX using python-docx.

        If the document is already in DOCX (application/vnd.openxmlformats-officedocument.wordprocessingml.document
        or application/msword), returns it unchanged.
        For other supported formats, first converts to plain text then creates a DOCX with the text.

        Args:
            content (bytes): Raw binary content of the document.
            mimetype (str): MIME type of the document.

        Returns:
            bytes: Converted DOCX content.
        """
        # 対象の MIME タイプ（DOCX／Word）の場合はそのまま返す
        docx_mimetypes = {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        }
        if mimetype in docx_mimetypes:
            return content

        # それ以外の場合は、まずテキストに変換する
        plain_text = self.convert_to_text(content, mimetype).decode("utf-8")

        # python-docx を利用して新規 DOCX を作成
        doc = Document()

        # テキストを行単位に分割し、各行を段落として追加する
        for line in plain_text.splitlines():
            doc.add_paragraph(line)

        # BytesIO に保存してバイナリデータを返す
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    # def _convert_epub_to_text(self, content: bytes) -> bytes:
    #     """Extracts text from EPUB using ebooklib."""

    #     book = epub.read_epub(BytesIO(content))
    #     text_content = ""
    #     for item in book.get_items():
    #         if item.get_type() == epub.EPUB_TEXT:
    #             text_content += item.get_content().decode("utf-8") + "\n"
    #     return text_content.encode("utf-8")

    # def _convert_docx_to_text(self, content: bytes) -> bytes:
    #     """Extracts text from DOCX using python-docx."""
    #     from docx import Document

    #     doc = Document(BytesIO(content))
    #     text_content = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    #     return text_content.encode("utf-8")


document_processor = DocumentProcessor()
